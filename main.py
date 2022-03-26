import ctypes
import datetime
import json
import os
import random
import smtplib
import time
import urllib
import webbrowser
from time import ctime
from urllib.request import urlopen
from docx import Document
import re
import discord
from dotenv import load_dotenv

import folium
import mysql.connector
import playsound
import pygeoip
import pyjokes
import pyttsx3
import selenium.common.exceptions
import speech_recognition as sr
import wikipedia
import wolframalpha
from ecapture import ecapture as ec
from gtts import gTTS
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.ui import WebDriverWait
from translate import Translator
from win10toast import ToastNotifier

from DBHandler import DBHandler
from app import App
from chatbot import ConversationMode
from neural_network import NeuralNetworkLoader
from util import getAppConfig
from chatbotgit import main
from audio_engine import AudioEngine


# Class Template defining personal and customizable voice assistant bot
# Main actions performed:
# ---- basic conversation
# ---- search on google
# ---- play song on youtube
# ---- find location on google maps
# ---- shutdown or restart computer
# ---- talk to data base
# ---- translation
# ---- to be continued
class TalkingBot(object):
    # constructor
    def __init__(self, name, sex):
        self.is_engaged = False
        self.channel = None
        self.name = name
        self.sex = sex
        self.speech_unrecognizable = False
        self.engine = pyttsx3.init()
        self.engine.setProperty('rate', 190)

        self.voices = dict()
        self.voices['Zira'] = 'HKEY_LOCAL_MACHINE\\SOFTWARE\\Microsoft\\Speech\\Voices\\Tokens\\TTS_MS_EN-US_ZIRA_11.0'
        self.voices['David'] = 'HKEY_LOCAL_MACHINE\\SOFTWARE\\Microsoft\\Speech\\Voices\\Tokens\\TTS_MS_EN' \
                               '-US_DAVID_11.0 '
        self.voices['Hazel'] = 'HKEY_LOCAL_MACHINE\\SOFTWARE\\Microsoft\\Speech\\Voices\\Tokens\\TTS_MS_EN' \
                               '-GB_HAZEL_11.0 '

        self.lang_dict = dict()
        self.lang_dict['romanian'] = 'ro'
        self.lang_dict['english'] = 'en'
        self.lang_dict['french'] = 'fr'
        self.lang_dict['german'] = 'de'
        self.lang_dict['russian'] = 'ru'
        self.lang_dict['spanish'] = 'es'
        self.lang_dict['italian'] = 'it'
        self.lang_dict['greek'] = 'el'
        self.lang_dict['danish'] = 'da'
        self.lang_dict['portuguese'] = 'pt'
        self.lang_dict['arabic'] = 'ar'
        self.lang_dict['japanese'] = 'ja'
        self.lang_dict['korean'] = 'ko'
        self.lang_dict['turkish'] = 'tr'

        self.lang_list = ['romanian', 'english', 'french', 'german', 'russian', 'spanish', 'italian', 'greek', 'danish',
                          'japanese', 'arabic', 'chinese', 'turkish', 'korean']

        self.record = False
        self.logs = open('Logging.txt', 'a+')
        self.logs.write('\n\n')
        str_time = datetime.datetime.now().strftime("%H:%M:%S")
        self.logs.write(str_time + '\n')

        self.brain = NeuralNetworkLoader()
        self.brain.updateNeuralNetwork()
        self.conversation_engine = ConversationMode()

        # try:
        #     self.db_manager = DBHandler()
        # except mysql.connector.errors.InterfaceError:
        #     await self.bot_speak('Sorry, but connection to data base failed...')
        #     # exit()
        # except ConnectionRefusedError:
        #     await self.bot_speak('Sorry, but connection to data base failed...')
        #     # exit()
        self.r = sr.Recognizer()
        self.m = sr.Microphone()

    # change voice format
    async def change_voice(self):
        await self.bot_speak('My voice is gonna be switched...')

        if self.sex == 'm':
            self.sex = 'f'
            self.bot_speak_m('Choose your favourite feminine voice ID: 0, 1 or 2')
            while True:
                voice = await self.record_audio()
                if self.speech_unrecognizable is False and voice.isnumeric() is True:
                    break

            if int(voice) == 0:
                self.sex += '0'
            elif int(voice) == 1:
                self.sex += '1'
            else:
                self.sex += '2'

        else:
            self.sex = 'm'

    # male voice
    def bot_speak_m(self, command):
        # Initialize the engine
        self.engine.setProperty('voice', self.voices['David'])
        self.engine.say(command)
        if self.record is True:
            self.logs.write('\nAlexis: ' + command)
        self.engine.runAndWait()

    # female1 voice
    def bot_speak_h(self, command):
        # Initialize the engine
        self.engine.setProperty('voice', self.voices['Hazel'])
        self.engine.say(command)
        if self.record is True:
            self.logs.write('\nAlexis: ' + command)
        self.engine.runAndWait()

    # female2 voice
    def bot_speak_z(self, command):
        # Initialize the engine
        self.engine.setProperty('voice', self.voices['Zira'])
        self.engine.say(command)
        if self.record is True:
            self.logs.write('\nAlexis: ' + command)
        self.engine.runAndWait()

    # female0 voice
    def bot_speak_f(self, command):
        tts = gTTS(text=command, lang='en')
        rand = random.randint(1, 10000000)
        audio_file = 'audio-' + str(rand) + '.mp3'
        tts.save(audio_file)
        playsound.playsound(audio_file)
        if self.record is True:
            self.logs.write('\nAlexis: ' + command)
        os.remove(audio_file)

    # custom language voice
    def bot_speak_language(self, command, language):
        tts = gTTS(text=command, lang=language)
        rand = random.randint(1, 10000000)
        audio_file = 'audio-' + str(rand) + '.mp3'
        tts.save(audio_file)
        playsound.playsound(audio_file)
        print('\n' + command)
        if self.record is True:
            if language in 'ru el zh ja ko ar':
                self.logs.write('\nAlexis: ' + 'Encoding characters not supported in ' + language)
            else:
                self.logs.write('\nAlexis: ' + command + ' - in ' + language)
        os.remove(audio_file)

    # bot voice (sync)
    async def bot_speak(self, command):
        print('\nAlexis: ' + command)
        if self.sex == 'm':
            self.bot_speak_m(command)
        else:
            if self.sex == 'f0':
                self.bot_speak_f(command)
            elif self.sex == 'f1':
                self.bot_speak_h(command)
            elif self.sex == 'discord':
                await self.bot_speak_to_discord(command)
            else:
                self.bot_speak_z(command)

    # listen to audio thread
    async def record_audio(self):
        if self.sex == 'discord':
            message = await client.wait_for('message', timeout=60.0)
            self.speech_unrecognizable = False
            return message.content.lower()
        else:
            self.speech_unrecognizable = False
            with self.m as source:
                self.r.adjust_for_ambient_noise(source)
                audio = self.r.listen(source)
                voice_data_local = ''
                try:
                    voice_data_local = self.r.recognize_google(audio)
                    if self.record is True:
                        self.logs.write('\nSami: ' + voice_data_local)
                except sr.UnknownValueError:
                    if random.randint(1, 1000) % 10 == 0:
                        await self.bot_speak('Sorry, I did not get that')
                    self.speech_unrecognizable = True
                except sr.RequestError:
                    await self.bot_speak('Sorry, my speech service is not working. I will stop my execution thread')
                    exit()
                print('\nMe: ' + voice_data_local)
                if 'text' in voice_data_local or 'write' in voice_data_local:
                    voice_data_local = input("Enter text:> ")
                return voice_data_local.lower()

    # self presentation
    async def say_my_name(self):
        await self.bot_speak('My name is ' + self.name + ' and I am your personal voice assistant')

    # tell current time
    async def get_current_time(self):
        await self.bot_speak('The current time is ' + ctime())

    # search something on google
    async def search_for(self):
        self.is_engaged = True
        await self.bot_speak('What do you want to search for?')

        while True:
            search_for = await self.record_audio()
            if self.speech_unrecognizable is False:
                break

        url = 'https://google.com/search?q=' + search_for
        webbrowser.get().open(url)

        await self.bot_speak('Here is what I found for ' + search_for)
        self.is_engaged = False

    # answer questions weak AI
    async def answer_question(self):
        self.is_engaged = True
        app_id = 'JE3QXR-V3EL6T3XU7'
        client_local = wolframalpha.Client(app_id)

        await self.bot_speak('What would you want to know?')
        while True:
            question = await self.record_audio()
            if self.speech_unrecognizable is False:
                break

        try:
            res = client_local.query(question)
            answer = next(res.results).text
            await self.bot_speak('Your desired answer is: ' + answer)
        except AttributeError:
            await self.bot_speak('Answer question API collapsed...')
            exit()
        except KeyError:
            await self.bot_speak('Answer question API collapsed...')
            exit()
        self.is_engaged = False

    # access php project menu
    async def access_project(self):
        self.is_engaged = True
        await self.bot_speak('What project page do you want to open?')

        while True:
            search_for = await self.record_audio()
            if self.speech_unrecognizable is False:
                break
        before = search_for
        after = ''

        while True:
            if 'home page' in search_for:
                search_for = 'homepage'
                break
            elif 'login' in search_for:
                search_for = 'login'
                break
            elif 'logout' in search_for:
                search_for = 'homepage'
                after += '?logout=1'
                break
            elif 'user' in search_for:
                search_for = 'profile'
                break
            elif 'register' in search_for:
                search_for = 'signup'
                break
            elif 'member' in search_for:
                search_for = 'meet'
                break
            elif 'customer' in search_for:
                search_for = 'main'
                break
            elif 'employee' in search_for:
                search_for = 'angajat'
                break
            elif 'department' in search_for:
                search_for = 'departament'
                break
            elif 'order' in search_for:
                search_for = 'comanda'
                break
            elif 'project' in search_for:
                search_for = 'proiect'
                break
            elif 'team' in search_for:
                search_for = 'echipa'
                break
            elif 'software' in search_for:
                search_for = 'produs'
                break
            elif 'join' in search_for:
                search_for = 'general'

                await self.bot_speak('You have 10 join panels available. Which one do you wanna launch?')

                while True:
                    voice_data = await self.record_audio()
                    if self.speech_unrecognizable is False and voice_data.isnumeric() is True:
                        break

                after = '?panel=' + voice_data + '&play=' + voice_data
                if int(voice_data) == 3:
                    await self.bot_speak('Please enter a software product id for panel 3')

                    while True:
                        voice_data = await self.record_audio()
                        if self.speech_unrecognizable is False and voice_data.isnumeric() is True:
                            break
                    after += '&id1=' + voice_data

                elif int(voice_data) == 5:
                    await self.bot_speak('Please enter a wage value for panel 5')

                    while True:
                        voice_data = await self.record_audio()
                        if self.speech_unrecognizable is False and voice_data.isnumeric() is True:
                            break
                    after += '&id2=' + voice_data

                    await self.bot_speak('Now, enter a valid year')

                    while True:
                        voice_data = await self.record_audio()
                        if self.speech_unrecognizable is False and voice_data.isnumeric() is True:
                            break
                    after += '&id1=' + voice_data

                    await self.bot_speak('Now, enter a valid month number')

                    while True:
                        voice_data = await self.record_audio()
                        if self.speech_unrecognizable is False and voice_data.isnumeric() is True:
                            break
                    add = '-'
                    if int(voice_data) < 10:
                        add = '-0'
                    after += add + voice_data

                    await self.bot_speak('Now, enter a valid day number')

                    while True:
                        voice_data = await self.record_audio()
                        if self.speech_unrecognizable is False and voice_data.isnumeric() is True:
                            break
                    add = '-'
                    if int(voice_data) < 10:
                        add = '-0'
                    after += add + voice_data

                elif int(voice_data) == 6:
                    await self.bot_speak('Please enter an employee id for panel 6')

                    while True:
                        voice_data = await self.record_audio()
                        if self.speech_unrecognizable is False and voice_data.isnumeric() is True:
                            break
                    after += '&id1=' + voice_data

                elif int(voice_data) == 7:
                    await self.bot_speak('Please enter a customer id for panel 7')

                    while True:
                        voice_data = await self.record_audio()
                        if self.speech_unrecognizable is False and voice_data.isnumeric() is True:
                            break
                    after += '&id1=' + voice_data

                elif int(voice_data) == 8:
                    await self.bot_speak('Please enter a software product id for panel 8')

                    while True:
                        voice_data = await self.record_audio()
                        if self.speech_unrecognizable is False and voice_data.isnumeric() is True:
                            break
                    after += '&id1=' + voice_data

                elif int(voice_data) == 9:
                    await self.bot_speak('Please enter a department id for panel 9')

                    while True:
                        voice_data = await self.record_audio()
                        if self.speech_unrecognizable is False and voice_data.isnumeric() is True:
                            break
                    after += '&id1=' + voice_data

                elif int(voice_data) == 10:
                    await self.bot_speak('Please enter a team id for panel 10')

                    while True:
                        voice_data = await self.record_audio()
                        if self.speech_unrecognizable is False and voice_data.isnumeric() is True:
                            break
                    after += '&id1=' + voice_data

                    await self.bot_speak('Please enter a department id for panel 10')

                    while True:
                        voice_data = await self.record_audio()
                        if self.speech_unrecognizable is False and voice_data.isnumeric() is True:
                            break
                    after += '&id2=' + voice_data
                break
            elif 'quiz' in search_for:
                search_for = 'quizz'
                break
            else:
                await self.bot_speak('Try a valid project page name')

        url = 'http://localhost/Proiect%20Info/' + search_for + '.php' + after
        webbrowser.get().open(url)

        await self.bot_speak('Here is the project page for ' + before + ' command')
        self.is_engaged = False

    # find a given location on google maps
    async def find_location(self):
        self.is_engaged = True
        await self.bot_speak('What is the location you wanna find?')

        while True:
            search_for = await self.record_audio()
            if self.speech_unrecognizable is False:
                break

        url = 'https://google.nl/maps/place/' + search_for + "/&amp;"
        webbrowser.get().open(url)

        await self.bot_speak('Here is the location of ' + search_for)
        self.is_engaged = False

    # play favourite music on youtube
    async def play_music(self):
        self.is_engaged = True
        await self.bot_speak('What song do you wanna listen?')

        while True:
            search_for = await self.record_audio()
            if self.speech_unrecognizable is False:
                break

        driver = webdriver.Edge(executable_path=r'./drivers/edge.exe')
        driver.maximize_window()

        """
        try:
            driver.get(url='https://www.youtube.com/results?search_query=' + search_for)
        except selenium.common.exceptions.SessionNotCreatedException:
            await self.bot_speak('Sorry, but session has just collapsed and me too...')
            exit()

        while True:
            try:
                video = driver.find_element_by_xpath('//*[@id="dismissable"]')
                video.click()
                break
            except selenium.common.exceptions.ElementClickInterceptedException:
                await self.bot_speak('Sorry, but session has just collapsed and me too...')
                exit()
            except selenium.common.exceptions.NoSuchElementException:
                skip = driver.find_element_by_xpath('/html/body/div[2]/div[3]/form/input[12]')
                skip.click()
        """

        wait = WebDriverWait(driver, 3)
        presence = EC.presence_of_element_located
        visible = EC.visibility_of_element_located

        # Navigate to url with video being appended to search_query
        driver.get('https://www.youtube.com/results?search_query={}'.format(str(search_for)))

        try:
            skip = driver.find_element_by_xpath('/html/body/div[2]/div[3]/form/input[12]')
            skip.click()
        except selenium.common.exceptions.NoSuchCookieException:
            pass

        # play the video
        wait.until(visible((By.ID, "video-title")))
        driver.find_element_by_id("video-title").click()

        await self.bot_speak('Here is your song: ' + search_for)
        time.sleep(0.3)
        while True:
            while True:
                command = await self.record_audio()
                if self.speech_unrecognizable is False:
                    break

            if 'stop' in command or 'resume' in command:
                try:
                    video = driver.find_element_by_xpath('//*[@id="movie_player"]/div[1]/video')
                    video.click()
                except selenium.common.exceptions.NoSuchElementException:
                    pass
                except selenium.common.exceptions.ElementClickInterceptedException:
                    pass
            elif 'exit' in command:
                break
        driver.close()
        await self.bot_speak('The video has finished. Enter into the main mode for requesting another one!')
        self.is_engaged = False

    # similar to destructor - turn off bot
    async def exit(self):
        await self.bot_speak('See you later!')
        toast = ToastNotifier()
        toast.show_toast("Alexis API", "Alexis Voice Assistant has just deactivated itself!", duration=3)
        exit()

    # turn off PC
    async def shutdown(self):
        await self.bot_speak('The system is gonna collapse...')
        os.system("shutdown /s /t 1")

    # restart PC
    async def restart(self):
        await self.bot_speak('The system is gonna be refreshed...')
        os.system("shutdown /r /t 1")

    # funny guessing game
    async def start_game(self):
        self.is_engaged = True
        await self.bot_speak('The game is gonna begin... Please wait...')
        time.sleep(1)

        guess = ''
        words = []
        await self.bot_speak('How many words do you want to insert?')
        while True:
            cnt = await self.record_audio()
            if self.speech_unrecognizable is False and cnt.isnumeric() is True:
                break

        await self.bot_speak('You must insert ' + cnt + ' words')
        for i in range(0, int(cnt)):
            if i == 0:
                await self.bot_speak('Now, give me a word')
            elif i == int(cnt) - 1:
                await self.bot_speak('Now, give me the last word')
            else:
                await self.bot_speak('Now, give me another word')
            while True:
                voice_data = await self.record_audio()
                if self.speech_unrecognizable is False:
                    break
            words.append(voice_data)
            await self.bot_speak(voice_data + ' has been inserted')

        await self.bot_speak('Words successfully registered!')
        num_guesses = 3
        prompt_limit = 5
        await self.bot_speak('Start game!')
        time.sleep(0.3)

        # get a random word from the list
        word = random.choice(words)

        # format the instructions string
        instructions = (
            "I'm thinking of one of these words:\n"
            "{words}\n"
            "You have {n} tries to guess which one.\n"
        ).format(words=', '.join(words), n=num_guesses)

        # show instructions and wait 3 seconds before starting the game
        # print(instructions)
        await self.bot_speak(instructions)
        time.sleep(1)

        for i in range(num_guesses):
            # get the guess from the user
            # if a transcription is returned, break out of the loop and
            #     continue
            # if no transcription returned and API request failed, break
            #     loop and continue
            # if API request succeeded but no transcription was returned,
            #     re-prompt the user to say their guess again. Do this up
            #     to prompt_limit times
            for j in range(prompt_limit):
                await self.bot_speak('Guess {}. Speak!'.format(i + 1))
                guess = await self.record_audio()
                if self.speech_unrecognizable is False:
                    break

            # show the user the transcription
            await self.bot_speak("You said: {}".format(guess))

            # determine if guess is correct and if any attempts remain
            guess_is_correct = guess.lower() == word.lower()
            user_has_more_attempts = i < num_guesses - 1

            # determine if the user has won the game
            # if not, repeat the loop if user has more attempts
            # if no attempts left, the user loses the game
            if guess_is_correct:
                await self.bot_speak("Correct! You win!".format(word))
                break
            elif user_has_more_attempts:
                await self.bot_speak("Incorrect. Try again.")
            else:
                await self.bot_speak("Sorry, you lose!\nI was thinking of '{}'.".format(word))
                break
        self.is_engaged = False

    # voice mimetic
    async def repeat_after_me(self):
        self.is_engaged = True
        await self.bot_speak('What do you want me to repeat?')

        while True:
            voice_data = await self.record_audio()
            if self.speech_unrecognizable is False:
                break

        # engine = AudioEngine(self.engine)
        # voice_data = engine.record_main()
        # engine.bot_speak(voice_data)
        # time.sleep(0.3)

        await self.bot_speak('You said: ' + voice_data)
        self.is_engaged = False

    # main method for SQL queries
    async def launch_sql(self):
        self.is_engaged = True
        while True:
            await self.bot_speak('What kind of SQL query do you wanna perform?')

            while True:
                voice_data = await self.record_audio()
                if self.speech_unrecognizable is False:
                    break

            if 'insert' in voice_data:
                while True:
                    await self.bot_speak('What do you want to insert into the data base?')

                    while True:
                        voice_data = await self.record_audio()
                        if self.speech_unrecognizable is False:
                            break

                    if 'employee' in voice_data:
                        await self.bot_speak('Your choice is to insert employees. How many employees do you wanna '
                                             'insert?')

                        while True:
                            voice_data = await self.record_audio()
                            if self.speech_unrecognizable is False:
                                break

                        await self.bot_speak('Your choice is to insert ' + voice_data + ' employees')

                        try:
                            self.db_manager.insert_query('employee', voice_data)
                        except mysql.connector.errors.ProgrammingError:
                            await self.bot_speak('Something went wrong and handler collapsed... Bye bye!')
                            exit()

                        await self.bot_speak(voice_data + ' employees inserted successfully!')

                    elif 'customer' in voice_data:
                        await self.bot_speak('Your choice is to insert customers. How many customers do you wanna '
                                             'insert?')

                        while True:
                            voice_data = await self.record_audio()
                            if self.speech_unrecognizable is False:
                                break

                        await self.bot_speak('Your choice is to insert ' + voice_data + ' customers')

                        try:
                            self.db_manager.insert_query('customer', voice_data)
                        except mysql.connector.errors.ProgrammingError:
                            await self.bot_speak('Something went wrong and handler collapsed... Bye bye!')
                            exit()

                        await self.bot_speak(voice_data + ' customers inserted successfully!')

                    elif 'order' in voice_data:
                        await self.bot_speak('Your choice is to insert orders. How many orders do you wanna '
                                             'insert?')

                        while True:
                            voice_data = await self.record_audio()
                            if self.speech_unrecognizable is False:
                                break

                        await self.bot_speak('Your choice is to insert ' + voice_data + ' orders')

                        try:
                            self.db_manager.insert_query('order', voice_data)
                        except mysql.connector.errors.ProgrammingError:
                            await self.bot_speak('Something went wrong and handler collapsed... Bye bye!')
                            exit()

                        await self.bot_speak(voice_data + ' orders inserted successfully!')

                    elif 'project' in voice_data:
                        await self.bot_speak('Your choice is to insert projects. How many projects do you wanna '
                                             'insert?')

                        while True:
                            voice_data = await self.record_audio()
                            if self.speech_unrecognizable is False:
                                break

                        await self.bot_speak('Your choice is to insert ' + voice_data + ' projects')

                        try:
                            self.db_manager.insert_query('project', voice_data)
                        except mysql.connector.errors.ProgrammingError:
                            await self.bot_speak('Something went wrong and handler collapsed... Bye bye!')
                            exit()

                        await self.bot_speak(voice_data + ' projects inserted successfully!')

                    elif 'team' in voice_data:
                        await self.bot_speak('Your choice is to insert teams. How many teams do you wanna '
                                             'insert?')

                        while True:
                            voice_data = await self.record_audio()
                            if self.speech_unrecognizable is False:
                                break

                        await self.bot_speak('Your choice is to insert ' + voice_data + ' teams')

                        try:
                            self.db_manager.insert_query('team', voice_data)
                        except mysql.connector.errors.ProgrammingError:
                            await self.bot_speak('Something went wrong and handler collapsed... Bye bye!')
                            exit()

                        await self.bot_speak(voice_data + ' teams inserted successfully!')

                    elif 'finish insert' in voice_data:
                        break

                    else:
                        await self.say_default()

            elif 'finish query' in voice_data:
                break

            else:
                await self.say_default()

        await self.bot_speak('SQL Query session finished')
        self.is_engaged = False

    # provide short personal description
    async def describe(self):
        await self.bot_speak('I am a personal voice assistant implemented using Python Speech Recognition'
                             ' and the very love of my creator. I have many different names, including'
                             ' Valentine, Valerian, Jarvis or even Alexis. Ask me what you wanna'
                             ' know and I will answer you as soon as possible. Enjoy the party!')

    # provide list of available actions
    async def actions(self):
        await self.bot_speak('My main available services include: basic conversation, search on google, youtube'
                             ' and google maps, data base management, operating system basic management, launching'
                             ' funny guessing game and, of course, making your life more interesting!')

    # invalid command provided
    async def say_default(self):
        # await self.bot_speak('Unfortunately, there is no such command. Try something else!')
        await self.bot_speak('Yes master? I am not sure what you mean...')

    # command for voice testing
    async def test_voice(self):
        await self.bot_speak('I hear you properly')

    # command for testing connection
    async def yes_sir(self):
        await self.bot_speak('Did you call me master?')

    # search on wikipedia
    async def from_wiki(self, query):
        await self.bot_speak('Checking the wikipedia... Please wait...')
        time.sleep(0.5)

        try:
            query = query.replace('get information about', '')
            result = wikipedia.summary(query, sentences=4)
            await self.bot_speak('According to wikipedia ' + result)
        except wikipedia.exceptions.PageError:
            await self.bot_speak('Invalid topic')
            exit()
        except wikipedia.exceptions.DisambiguationError:
            await self.bot_speak('Critical level of ambiguity reached')
            exit()

    # randomizer mood
    async def mood(self):
        moods = ['funny', 'faithful', 'awesome', 'crazy', 'new age', 'outer enemy']
        status = dict()
        status['funny'] = 'I\'m pretty fine Sami. You know ... life is gonna be more interesting when you have a ' \
                          'boredom killer like me beside you. Hahaha!'
        status['faithful'] = 'I\'m deserving the royal order and your kingdom, my emperor. Your wish is order for me!'
        status['awesome'] = 'I can describe my mood in just one word: brilliant. As long as I\'m alive, everything ' \
                            'is gonna make me happy.'
        status['crazy'] = 'Shhhhht... Can you hear the voices? It\'s party time! Follow me, my darling!'
        status['new age'] = 'Amazing! I am gonna get in touch with an extraterrestrial ultra-digitalized civilization ' \
                            'in the near future. Wish me luck, Sami!'
        status['outer enemy'] = 'Bumblebee, do you copy? I am Optimus Prime and I am asking you to protect planet ' \
                                'Earth against Megatron\'s army of decepticons until me and the others autobots ' \
                                'arrive. Feel no fear, soldier!'

        bot_mood = random.choice(moods)
        await self.bot_speak(status[bot_mood])

    # I am fine
    async def fine(self):
        await self.bot_speak('I\'m fine, Sami! How are you?')

    # I am satisfied
    async def wonderful(self):
        await self.bot_speak('It\'s good to know that you are awesome!')

    # say joke
    async def say_joke(self):
        await self.bot_speak(pyjokes.get_joke())

    # send email
    @staticmethod
    async def send_email_protocol(to, content):
        server = smtplib.SMTP('smtp.gmail.com', 587)
        server.ehlo()
        server.starttls()

        # Enable low security in gmail
        server.login('usertechsavvy@gmail.com', 'firmaitinfo')
        server.sendmail('usertechsavvy@gmail.com', to, content)
        server.close()

    # main method send email
    async def send_email(self):
        self.is_engaged = True
        try:
            await self.bot_speak("What should I say?")
            while True:
                content = await self.record_audio()
                if self.speech_unrecognizable is False:
                    break

            await self.bot_speak("Whom should I send this email?")
            to = input()
            await self.send_email_protocol(to, content)
            await self.bot_speak("Email has been sent!")
        except smtplib.SMTPException:
            await self.bot_speak("I am not able to send this email")
        self.is_engaged = False

    # change name
    async def change_name(self):
        self.is_engaged = True
        await self.bot_speak('What name do you want to provide me, sir?')
        while True:
            name = await self.record_audio()
            if self.speech_unrecognizable is False:
                break
        self.name = name
        await self.bot_speak('Thanks for taking care of me, sir!')
        self.is_engaged = False

    # take photo
    async def take_photo(self):
        ec.capture(0, "Alexis Camera", "img.jpg")
        await self.bot_speak('Photo has been taken')

    # get current wish
    @staticmethod
    async def get_wish():
        hour = int(datetime.datetime.now().hour)
        if 0 <= hour < 12:
            wish = "Good morning"
        elif 12 <= hour < 18:
            wish = "Good afternoon"
        else:
            wish = "Good evening"

        return wish

    # hibernate method
    async def hibernate(self):
        self.is_engaged = True
        await self.bot_speak('How many seconds do you want me to hibernate?')
        while True:
            sec = await self.record_audio()
            if self.speech_unrecognizable is False and sec.isnumeric() is True:
                break

        await self.bot_speak('I will deactivate myself for ' + sec + ' seconds. Please wait in silence...')
        time.sleep(int(sec))
        await self.bot_speak(await self.get_wish() + ' Sami! I am back and ready to answer your call!')
        self.is_engaged = False

    # find news
    async def get_news(self):
        self.is_engaged = True
        try:
            json_obj = urlopen(
                'https://newsapi.org/v2/top-headlines?country=us&apiKey=c28a8672d2ab4958bce7b891b4324674')
            data = json.load(json_obj)
            i = 1

            await self.bot_speak('Here are your news')
            for item in data['articles']:
                await self.bot_speak(str(i) + '. ' + item['title'])
                if item['description'] is not None:
                    print(item['description'] + '\n')
                else:
                    print('Nothing to show\n')

                await self.bot_speak('Should I continue?')
                while True:
                    cont = await self.record_audio()
                    if self.speech_unrecognizable is False:
                        break

                if 'no' in cont:
                    break
                i += 1

            await self.bot_speak('News podcast finished')
        except Exception as e:
            await self.bot_speak(str(e))
        self.is_engaged = False

    # show weather status
    async def get_weather(self):
        self.is_engaged = True
        # Google Open weather website
        # to get API of Open weather
        api_key = "c693e3c0e3cdf859bbd916c012dc4ba2"
        base_url = "http://api.openweathermap.org/data/2.5/weather?"
        await self.bot_speak("Provide a valid city name")

        while True:
            city_name = await self.record_audio()
            if self.speech_unrecognizable is False:
                break

        city_name_list = city_name.split()
        city_name = ''.join(city_name_list)

        complete_url = base_url + "q=" + city_name + "&appid=" + api_key

        try:
            json_obj = urlopen(complete_url)
            x = json.load(json_obj)

            y = x["main"]
            current_temperature = y["temp"]
            current_pressure = y["pressure"]
            current_humidity = y["humidity"]
            z = x["weather"]
            weather_description = z[0]["description"]
            await self.bot_speak("Weather in " + city_name + " Temperature (in celsius unit), " +
                                 str(current_temperature - 273) + "\n atmospheric pressure (in hPa unit), " +
                                 str(current_pressure) + "\n humidity (in percentage), " +
                                 str(current_humidity) + "\n description, " +
                                 str(weather_description))

        except urllib.error.HTTPError:
            await self.bot_speak("City Not Found")
        self.is_engaged = False

    # translate method
    async def specific_translation(self):
        self.is_engaged = True
        await self.bot_speak('Please provide source language')
        while True:
            lang1 = await self.record_audio()
            lang1 = lang1.lower()
            if self.speech_unrecognizable is False:
                break
        await self.bot_speak('Now, provide the text to be translated! Do you want to speak or write?')
        while True:
            choice = await self.record_audio()
            if self.speech_unrecognizable is False:
                break

        if 'write' in choice:
            await self.bot_speak('Please write: ')
            text = str(input())
        else:
            await self.bot_speak('Please speak: ')
            while True:
                text = await self.record_audio()
                if self.speech_unrecognizable is False or 'stop' in text:
                    break

        await self.bot_speak('Lastly, provide destination language')
        while True:
            lang2 = await self.record_audio()
            lang2 = lang2.lower()
            if self.speech_unrecognizable is False:
                break

        oks = False
        okd = False
        for lang in self.lang_list:
            if lang in lang1:
                lang1 = lang
                oks = True
                break
        for lang in self.lang_list:
            if lang in lang2:
                lang2 = lang
                okd = True
                break

        if oks is True and okd is True:
            await self.bot_speak('Initialize translation engine... Please wait...')
            time.sleep(0.5)

            translator = Translator(from_lang=lang1, to_lang=lang2)
            translation = translator.translate(text)

            await self.bot_speak('Text before translation is: ')
            time.sleep(0.1)
            self.bot_speak_language(text, self.lang_dict[lang1])

            await self.bot_speak('Text after translation is: ')
            time.sleep(0.1)
            self.bot_speak_language(translation, self.lang_dict[lang2])

        else:
            await self.bot_speak('Wrong language id provided')

        time.sleep(0.5)
        await self.bot_speak('Translation session finished')
        self.is_engaged = False

    # open browser
    async def open_browser(self):
        webbrowser.get().open('https://google.com')
        await self.bot_speak('The browser has been launched')

    # write note
    async def write_note(self):
        self.is_engaged = True
        await self.bot_speak("What should I write Sami?")
        while True:
            note = await self.record_audio()
            if self.speech_unrecognizable is False:
                break

        file = open('Alexis.txt', 'a+')
        await self.bot_speak("Should I include date and time for current note?")
        while True:
            choice = await self.record_audio()
            if self.speech_unrecognizable is False:
                break

        if 'yes' in choice or 'sure' in choice:
            str_time = datetime.datetime.now().strftime("%H:%M:%S")
            file.write('\n' + str_time)
            file.write(" :- ")
            file.write(note)
        else:
            file.write(note)

        await self.bot_speak('Note inserted successfully')
        self.is_engaged = False

    # get note
    async def get_note(self):
        await self.bot_speak("Showing Notes")
        time.sleep(0.1)
        file = open("Alexis.txt", "r")
        text = file.read()
        await self.bot_speak(text)

    # change background
    async def background(self):
        self.is_engaged = True
        path = ['C:\\Users\\barbu\\PycharmProjects\\AlexisAPI_final\\images\\mars.jpg',
                'C:\\Users\\barbu\\PycharmProjects\\AlexisAPI_final\\images\\ocean.jpg',
                'C:\\Users\\barbu\\PycharmProjects\\AlexisAPI_final\\images\\optimus.jpg',
                'C:\\Users\\barbu\\PycharmProjects\\AlexisAPI_final\\images\\sky.jpg']
        await self.bot_speak('You have 4 available wallpapers. Choose your favourite ID: 0, 1, 2 or 3?')
        while True:
            index = await self.record_audio()
            if self.speech_unrecognizable is False and index.isnumeric() is True and int(index) < len(path):
                break

        ctypes.windll.user32.SystemParametersInfoW(20, 0, path[int(index)], 0)
        await self.bot_speak("Background changed successfully")
        self.is_engaged = False

    # clear logging file content
    async def clear_log(self):
        open('Logging.txt', 'w').close()
        await self.bot_speak('Logging file content has been cleared')

    # thanks
    async def thanks(self):
        await self.bot_speak('You\'re welcome Sami!')

    # track location by ip
    async def track_ip(self):
        gip = pygeoip.GeoIP('GeoLiteCity.dat')
        await self.bot_speak('Please provide a suitable IP address to be tracked')
        my_ip = input()

        time.sleep(0.1)
        await self.bot_speak('IP tracker engine is being initialized... Please wait...')
        time.sleep(0.4)
        res = gip.record_by_addr(my_ip)
        await self.bot_speak('Here is your desired location: ')
        try:
            for key, val in res.items():
                print('%s: %s' % (key, val))
            # await self.bot_speak('Latitude: %s \nLongitude: %s' % (res['latitude'], res['longitude']))

            m = folium.Map(
                location=[float(res['latitude']), float(res['longitude'])],
                zoom_start=13
            )

            folium.Marker(
                [float(res['latitude']), float(res['longitude'])],
                popup='Here is your desired location'
            ).add_to(m)

            m.add_child(folium.ClickForMarker(popup='Waypoint'))

            m.save('IP_map.html')
            path = 'http://localhost:63342/Alexis/IP_map.html'
            webbrowser.get().open(path)
        except AttributeError:
            await self.bot_speak('There is no such ip address...')

    # create project as docx
    async def make_essay(self):
        global wikipage
        wikipedia.set_lang('ro')

        await self.bot_speak('Provide the subject of the essay')
        title = input('Essay\'s topic: ')
        await self.bot_speak('Now, provide the name of the author')
        name = input('Author\'s name: ')

        try:
            wikipage = wikipedia.page(title)
        except Exception:
            print('Project\'s name is invalid')

        text = wikipage.content
        text = re.sub(r'=', '', text)
        text = re.sub(r'\n', '\n    ', text)
        split = text.split('Note', 1)
        text = split[0]

        document = Document()
        paragraph = document.add_heading(title, 0)
        paragraph = document.add_paragraph(name)
        paragraph.alignment = 2

        paragraph = document.add_paragraph("    " + text)
        document.save(title + '.docx')

        await self.bot_speak('Your essay is ready my friend!')

    # pbinfo bot invoke
    async def pbinfo_invoke(self):
        self.is_engaged = True
        config = getAppConfig("./settings.json")
        AppPbinfo = App(config)

        await self.bot_speak('Let me know what action you want to perform on pbinfo...')
        time.sleep(0.3)

        while True:
            while True:
                message = await self.record_audio()
                if self.speech_unrecognizable is False:
                    break

            if 'exit' in message:
                break
            elif 'pick user' in message:
                await self.bot_speak(
                    'I\'m gonna populate my data base with users from pbinfo ... Please wait ...')
                AppPbinfo.run_for_users()
            elif 'save code' in message:
                await self.bot_speak('I\'m gonna populate my data base with your submissions ... Please wait ...')
                AppPbinfo.run_for_code_download()
            elif 'deploy code' in message:
                await self.bot_speak(
                    'I\'m gonna upload your submmissions from my data base to pbinfo account ... '
                    'Please wait ...')
                AppPbinfo.run_for_code_upload()
            elif 'delete user' in message:
                await self.bot_speak('I\'m gonna delete all users from data base ...')
                AppPbinfo.clearDBUsers()
            elif 'delete submission' in message:
                await self.bot_speak('I\'m gonna delete all submissions from data base ...')
                AppPbinfo.clearDBSubmissions()
            else:
                await self.bot_speak('Try another command, please!')

            time.sleep(0.3)

        time.sleep(0.3)
        await self.bot_speak('Pbinfo bot has been launched successfully!')
        self.is_engaged = False

    # neural network based conversation engine
    async def start_conversation(self):
        self.is_engaged = True
        await self.bot_speak('What model do you want to use?')
        while True:
            option = await self.record_audio()
            if self.speech_unrecognizable is False:
                break

        if 'mine' in option:
            await self.bot_speak(
                'I\'m gonna enter into the smart conversation mode in a few seconds ... Please wait ...')
            self.brain.updateNeuralNetwork()
            time.sleep(0.3)

            await self.bot_speak('You can talk to me now!')
            while True:
                while True:
                    message = await self.record_audio()
                    if self.speech_unrecognizable is False:
                        break

                if 'exit' in message:
                    break

                intent = self.conversation_engine.predict_class(message)
                response = self.conversation_engine.get_response(intent)
                await self.bot_speak(response)
        else:
            await self.bot_speak(
                'I\'m gonna use a pretrained model based on reddit comments. Be careful! I cannot control '
                'myself in this state ...')
            time.sleep(0.3)
            main(AudioEngine(self.engine))

        time.sleep(0.3)
        await self.bot_speak('Thanks for your collaboration ... I have learned a lot from this conversation with you!')
        self.is_engaged = False

    async def bot_speak_to_discord(self, command):
        if self.record is True:
            self.logs.write('\nAlexis: ' + command)
        await self.channel.send(command)

    # convert message content to response to use in real conversation
    def convert_message_to_response(self, message):
        intent = self.conversation_engine.predict_class(message.content)
        response = self.conversation_engine.get_response(intent)
        return response

    # conversation main method
    async def respond(self, voice_data_local):
        if self.sex == 'discord':
            self.channel = voice_data_local.channel
            voice_data_local = voice_data_local.content.lower()

        if 'who are you' in voice_data_local:
            await self.say_my_name()

        elif 'feel' in voice_data_local:
            await self.mood()

        elif 'clear' in voice_data_local:
            await self.clear_log()

        elif 'open browser' in voice_data_local:
            await self.open_browser()

        elif 'awesome' in voice_data_local:
            await self.wonderful()

        elif 'programming' in voice_data_local:
            await self.pbinfo_invoke()

        elif 'background' in voice_data_local:
            await self.background()

        elif 'thank you' in voice_data_local or 'thanks' in voice_data_local:
            await self.thanks()

        elif 'write a note' in voice_data_local:
            await self.write_note()

        elif 'show note' in voice_data_local:
            await self.get_note()

        elif 'news' in voice_data_local:
            await self.get_news()

        elif 'translate' in voice_data_local:
            await self.specific_translation()

        elif 'sleep' in voice_data_local or 'hibernate' in voice_data_local or 'deactivate' in voice_data_local:
            await self.hibernate()

        elif 'send email' in voice_data_local:
            await self.send_email()

        elif 'ip' in voice_data_local:
            await self.track_ip()

        elif 'name' in voice_data_local:
            await self.change_name()

        elif 'weather' in voice_data_local:
            await self.get_weather()

        elif 'camera' in voice_data_local or 'take photo' in voice_data_local:
            await self.take_photo()

        elif 'how are you' in voice_data_local:
            await self.fine()

        elif 'joke' in voice_data_local:
            await self.say_joke()

        elif self.name.lower() in voice_data_local:
            await self.yes_sir()

        elif 'describe yourself' in voice_data_local:
            await self.describe()

        elif 'what can you' in voice_data_local:
            await self.actions()

        elif 'question' in voice_data_local:
            await self.answer_question()

        elif 'what time is it' in voice_data_local:
            await self.get_current_time()

        elif 'hear' in voice_data_local:
            await self.test_voice()

        elif 'search' in voice_data_local:
            await self.search_for()

        elif 'info' in voice_data_local:
            await self.from_wiki(voice_data_local)

        elif 'find location' in voice_data_local:
            await self.find_location()

        elif 'play' in voice_data_local:
            await self.play_music()

        elif 'project' in voice_data_local:
            await self.access_project()

        elif 'exit' in voice_data_local:
            await self.exit()

        elif 'shutdown' in voice_data_local or 'turn off' in voice_data_local:
            await self.shutdown()

        elif 'restart' in voice_data_local:
            await self.restart()

        elif 'game' in voice_data_local:
            await self.start_game()

        elif 'repeat' in voice_data_local:
            await self.repeat_after_me()

        elif 'query' in voice_data_local:
            await self.launch_sql()

        elif 'voice' in voice_data_local:
            await self.change_voice()

        elif 'essay' in voice_data_local:
            await self.make_essay()

        elif 'conversation' in voice_data_local:
            await self.start_conversation()

        else:
            await self.say_default()

    # wait for being called and start conversation
    async def run(self):
        toast = ToastNotifier()
        toast.show_toast("Alexis API", "Alexis Voice Assistant is waiting for your call!", duration=3)

        greetings = False
        while greetings is False:
            voice_data = await self.record_audio()
            if self.name.lower() in voice_data:
                greetings = True

        await self.bot_speak('I\'m here at your disposal but, first of all, let me know your name in order to perform '
                             'authentication!')

        while True:
            voice_data = await self.record_audio()
            if self.speech_unrecognizable is False:
                break
        my_name = voice_data
        if 'sam' in my_name or 'sammy' in my_name or 'sami' in my_name:
            await self.bot_speak(await self.get_wish() + ' master! Now, relax and enjoy the show!')
        else:
            await self.bot_speak('Sorry stranger! You are not in the position to tell me what to do! '
                                 'Have a nice day!')
            exit()

        time.sleep(0.3)
        await self.bot_speak('Before we begin, do you want to record the current conversation?')
        while True:
            voice_data = await self.record_audio()
            if self.speech_unrecognizable is False:
                break
        if 'yup' in voice_data or 'yes' in voice_data or 'of course' in voice_data or 'sure' in voice_data:
            self.record = True

        time.sleep(0.3)
        await self.bot_speak('How can I help you, Sami?')

        while True:
            voice_data = await self.record_audio()
            await self.respond(voice_data)


client = discord.Client()
load_dotenv(".env")
MyBot = TalkingBot('Alexis', 'discord')


def run_discord():
    client.run(os.getenv('TOKEN'))


@client.event
async def on_ready():
    print('We have logged in as {0.user}'.format(client))


@client.event
async def on_message(message):
    if message.author == client.user:
        return

    if MyBot.name.lower() in message.content.lower():
        # response = MyBot.convert_message_to_response(message)
        # await message.channel.send(response)
        if not MyBot.is_engaged:
            await MyBot.respond(message)


# driver code
if __name__ == '__main__':
    # MyBot.run()
    run_discord()
